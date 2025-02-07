import re
from pdfminer.high_level import extract_text
import fitz
import camelot
import pandas as pd
import os
from docx import Document


def GetPDFName(path):       # Get only the PDF file name from the whole path
    # This variable is used to map to the last character in the path combined with the pdf name
    index = len(path) - 1

    #Initialize a string variable named pdf_name to store the PDF file name without the whole path
    pdf_name = ''
    while(True):
        if path[index] == '/':
            break
        pdf_name += f'{path[index]}'
        index -= 1
    # Reverse text to be in the right order
    pdf_name = pdf_name[::-1]

    # Folder is named as the pdf file but without the last 4 characters '.pdf'
    pdf_name = pdf_name[:-4]
    return pdf_name



def IndexRepeatedArrayElements(array):
    for index in range(len(array)):
        counter = 1
        repeated = []
        repeated_count = 0
        for subindex in range(len(array) - index):
            if array[index + subindex] == array[index]:
                repeated.append(index + subindex)
                repeated_count += 1
        

        if (repeated_count > 1):
            for idx in repeated :
                array[idx] = array[idx] + f'_{counter}'
                counter += 1 
    return array
        

#################################
######## Extracting Text ########
#################################
def ExtractText(pdf_file):
    text = extract_text(pdf_file)
    return text



def GetTitle(text):
    title = ''
    for i in text:
        if (i != '\n'):
            title += f'{i}'
        else:
            break
    return title



def GetTextBeforeColon(full_text, index):
    # Decrement index to find elements before the colon
    index -= 1

    # Initialize a string variable named text to store text before colon in it
    text = ''
    if full_text[index] == ' ':
        index -= 1
    while(True):
        if full_text[index] == '\n':
            break
        text += f'{full_text[index]}'
        index -= 1
    # Reverse text to be in the right order
    text = text[::-1]
    return text



def GetElementsAfterColon(full_text, index):
    # Increment index to find elements after the colon
    index += 1

    # While the character is a white space, increment the index until we reach the values.
    while (full_text[index] == ' '):
        index += 1
    
    # Initialize a variable line as empty string to store characters of the line in it later 
    line = ''

    # While line is not ended
    while(full_text[index] != '\n'):
        line += f'{full_text[index]}'
        index += 1

    # Check the words that ends with the delimiters comma ',' or hyphen '-'
    pattern = re.compile(r'(\w+)[,-](?=\W|$)')

    # Find all matches in line
    elements = re.findall(pattern, line)

    '''
        NOTE:
            All words ending with comma or hyphen are included but the last element of the list is not added
    '''

    last_element = ''

    # Use the variable index now as the index for the last character of line
    index = len(line) - 1
    while((line[index] != ',') and (line[index] != '-')):
        last_element += f'{line[index]}'
        index -= 1
    
    # Reverse last_element to be in the right order
    last_element = last_element[::-1]


    # Check if last character is fullstop
    if last_element[-1] == '.':
        # Remove the fullstop
        last_element = last_element[:-1]

    # Add last_element to elements
    elements.append(last_element)

    # Return the list of elements
    return elements



def FindListsInText(text):
    list_names = []
    list_elements = []
    lists = {}
    for i in range(len(text)):
        if (text[i] == ':' and ((text[i+1] != '\n') or (text[i+1] == ' ' and full_text[i+2] != '\n'))):
            
            # Store name of current list in current_name
            current_name = GetTextBeforeColon(text, i)

            # Add current_name to list_names
            list_names.append(current_name)        

            # Store elements of current list in current_elements
            current_elements = GetElementsAfterColon(text, i)
        
            # Add current_elements to list_elements
            list_elements.append(current_elements)

    # Before adding the names of lists into a dictionary as keys with the elements as values, make each key unique. 
    list_names = IndexRepeatedArrayElements(list_names)

    for counter in range(len(list_names)):
        lists.update({list_names[counter] : list_elements[counter]})

    return lists


###################################
######## Extracting Tables ########
###################################

def ExtractTables(pdf_file):
    tables = camelot.read_pdf(pdf_file, pages='1-end')

    print(f"Found {tables.n} {"tables" if tables.n > 1 else "table"}")
    # If there is at least one table
    if tables.n != 0:
        # Get the PDF file name without the whole path
        folder = GetPDFName(pdf_file)

        if not os.path.exists(folder):
            os.mkdir(f'{folder}')

        output_folder = f"{folder}/tables"
        # Create the "tables" folder if it doesn't exist
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            print(f"[+] Created output folder: {output_folder}")

        # Export tables as CSV files in folder tables            
        tables.export(f'{output_folder}/tables.csv', f='csv')

    return tables




###################################
######## Extracting Images ########
###################################

def ExtractImages(pdf_file):
    pdf_file_path = pdf_file
    
    pdf_file = fitz.open(pdf_file)
    processed_xrefs = set()  # To store unique image XREFs

    for page_index in range(len(pdf_file)):
        page = pdf_file[page_index]
        image_list = page.get_images()  # get images on the page

        # printing number of images found in this page
        if image_list:
            print(f"Found {len(image_list)} {"images" if len(image_list) > 1 else "image"} on page {page_index + 1}")

        for image_index, img in enumerate(image_list, start=1):
            # get the XREF of the image
            xref = img[0]

            # Check if we've already processed this image
            if xref in processed_xrefs:
                print(f"[!] Skipping duplicate image with XREF {xref} on page {page_index + 1}")
                continue

            # Add the XREF to the set of processed images
            processed_xrefs.add(xref)

            # extract the image bytes
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]
            
            # get the image extension
            image_ext = base_image["ext"]

            # save the image in the "img" folder
            image_name = f"image{page_index+1}_{image_index}.{image_ext}"
            
            
            folder = GetPDFName(pdf_file_path)

            if not os.path.exists(folder):
                os.mkdir(f'{folder}')


            output_folder = f"{folder}/img"
            # Create the "img" folder if it doesn't exist
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                print(f"[+] Created output folder: {output_folder}")

            
            image_path = os.path.join(output_folder, image_name)  # Full path to save the image
            with open(image_path, "wb") as image_file:
                image_file.write(image_bytes)







def GetExtractedData(pdf_file, title, lists):
    document = Document()

    document.add_heading('Title', level = 1)
    title_p = document.add_paragraph(f'The title of "{pdf_file}" is: ')

    title_p.add_run(f'{title}').bold = True

    document.add_heading('Lists', level = 1)
    title_p = document.add_paragraph(f'The lists found inside "{pdf_file}" are:')

    for key, value in lists.items():
        document.add_paragraph(f'{key} : {value}', style='List Bullet')
    
    # Folder is named as the pdf file but without the extension
    folder = pdf_file[:-4]
    if not os.path.exists(folder):
        os.mkdir(f'{folder}')
    document.save(f'{folder}/Extracted_Data.docx')





print("Please, enter the path of PDF files you want to extract!")
print("If the path is the current path, press Enter")
path = input('If the path is the current path, just press Enter\n')
if path == '':
    path = os.getcwd()


for i in os.listdir(path):
    if i.endswith(".pdf"):
        print(f"Extracting {i} ...")

        pdf_file = f'{path}/{i}'
        full_text = ExtractText(pdf_file)

        title = GetTitle(full_text)

        lists = FindListsInText(full_text)

        ExtractTables(pdf_file)
        ExtractImages(pdf_file)

        GetExtractedData(i, title, lists)





